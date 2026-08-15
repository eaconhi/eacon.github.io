from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tmp" / "fonttools"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Inches, Pt, RGBColor
from fontTools.ttLib import TTFont as FontToolsTTFont
from lxml import etree
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


PUBLIC = ROOT / "public"
FONT_DIR = Path("/System/Library/Fonts/Supplemental")

INK = colors.HexColor("#111111")
MUTED = colors.HexColor("#505763")
VIOLET = colors.HexColor("#6F4BE8")
PALE = colors.HexColor("#F4F5F8")
LINE = colors.HexColor("#D9DCE3")

ZH_JOBS = [
    (
        "NAAI / EAE 国际学术与文化项目运营",
        "NAAI 核心管理层 / 国际项目负责人；EAE 网站内容运营 | 2025.03 - 至今",
        [
            "建立 NAAI / EAE 会员资料核验字段，识别身份、机构、职务、研究方向、代表成果和公开来源中的缺失、冲突与待确认项。",
            "运营 EAE Member、News 等网站内容，完成材料整理、英文书面核对、称谓与来源检查及发布后页面复核。",
            "统筹 Astria Awards 等影视文化项目的提名、影片与人员材料，跟进关键缺口、导演或评委沟通、证书文本和节点归档。",
            "为 AI Agent 设定来源优先级、禁止猜测项、输出格式和验收要求，复核事实与公开结果，并沉淀 Skill 和检查清单。",
        ],
    ),
    (
        "硕途教育科技集团",
        "常务总经理 | 2023 - 2025",
        [
            "汇总用户、内容、社群和合作方反馈，判断优先级并明确责任环节与处理节点。",
            "将资料缺口与流程卡点整理为可跟进事项，协调补充信息并推动解决。",
            "复核处理结果，根据反馈调整内容、服务与后续安排。",
        ],
    ),
    (
        "思源汉客文化传播有限公司",
        "内容运营负责人 / 新媒体运营 | 2018.03 - 2023.08",
        [
            "负责教育类内容的选题、编辑、发布和日常维护。",
            "根据用户反馈与内容表现调整选题、表达和发布节奏。",
            "组织线上活动与推广内容，跟进发布、互动和复盘，优化后续执行。",
        ],
    ),
    (
        "奎屯市第一中学",
        "初中英语教师 | 2016.09 - 2018.02",
        [
            "根据教学目标和学生差异拆解知识点、组织课堂活动。",
            "记录学习情况并沟通反馈，将问题转成教学调整与跟进动作。",
        ],
    ),
]

EN_JOBS = [
    (
        "NAAI & EAE / Academic and Cultural Project Operations",
        "NAAI Core Management Team / International Project Lead; EAE Website Content Operations | Mar 2025 - Present",
        [
            "Builds verification fields for NAAI and EAE member data, identifying gaps, conflicts and open questions across identity, institution, role, research area, representative work and public sources.",
            "Runs EAE Member, News and related website content through material organization, written-English review, title and source checks, and post-publication page review.",
            "Coordinates nomination, film and participant materials for Astria Awards and related programs, following critical gaps, director or jury communication, certificate wording and milestone records.",
            "Sets source priorities, no-guess boundaries, output formats and acceptance criteria for AI agents, reviews facts and public results, and captures reusable skills and checklists.",
        ],
    ),
    (
        "Shuotu Education Technology Group",
        "Executive General Manager | 2023 - 2025",
        [
            "Consolidated user, content, community and partner feedback, setting priorities, owners and action milestones.",
            "Turned missing materials and process blockers into trackable actions, coordinating information recovery and resolution.",
            "Reviewed outcomes and adjusted content, service and follow-up plans based on feedback.",
        ],
    ),
    (
        "Siyuan Hanke Culture Communication Co., Ltd.",
        "Content Operations Lead / New Media Operations | Mar 2018 - Aug 2023",
        [
            "Managed education content across topic selection, editing, publishing and ongoing maintenance.",
            "Adjusted topics, framing and publishing schedules based on feedback and content performance.",
            "Organized campaigns and promotional content, following publication, interaction and review to improve later execution.",
        ],
    ),
    (
        "Kuitun No.1 Middle School",
        "Junior High School English Teacher | Sep 2016 - Feb 2018",
        [
            "Structured concepts and classroom activities around learning goals and student readiness.",
            "Tracked progress and turned student and parent feedback into teaching adjustments and follow-up actions.",
        ],
    ),
]

ZH_CASES = [
    (
        "会员信息治理与内容质量闭环",
        [
            "<b>场景：</b>资料可能出现字段缺失、职务口径不统一、公开来源不清或页面显示不一致。",
            "<b>我的动作：</b>把业务要求转成核验字段与判断边界，借助 AI 提取信息，再人工复核身份、机构、专业背景、官方来源和页面状态，将问题转成具体修正动作。",
            "<b>交付：</b>会员资料核验记录、字段标准与修改清单、会员或新闻页面，以及发布检查流程与复核结果。",
        ],
    ),
    (
        "AI Agent 协作的运营工作流",
        [
            "<b>场景：</b>学术组织运营涉及公开来源、会员资料、新闻内容、页面状态和多个时间节点，任务反复且判断边界、验收口径容易不一致。",
            "<b>我的动作：</b>把工作拆成检索、字段整理、内容生成、页面检查和公开结果复核，为 AI Agent 设定来源优先级、输出格式、禁止猜测项和验收标准，并根据偏差持续修正规则。",
            "<b>交付：</b>结构化核验记录与修改清单、会员与新闻内容、来源与页面复核结果，以及可复用 Skill、工作流和检查清单。",
        ],
    ),
    (
        "影视文化奖项材料与节点协同",
        [
            "<b>场景：</b>奖项项目涉及提名、影片材料、导演或评委沟通、评审节点、证书文本和资料归档。",
            "<b>我的动作：</b>整理提名与影片材料，记录缺口和待确认项，跟进沟通，核对证书文字并完成归档。",
            "<b>交付：</b>提名与影片材料、待确认事项记录、沟通跟进信息、证书文本与归档资料。",
        ],
    ),
]

EN_CASES = [
    (
        "Member Information Governance & Content Quality",
        [
            "<b>Context:</b> Materials may include missing fields, inconsistent titles, unclear public sources or page-level discrepancies.",
            "<b>Actions:</b> Translates the operating need into verification fields and decision boundaries, uses AI for extraction, manually reviews identity, institution, professional background, public sources and page content, and turns issues into corrective actions.",
            "<b>Deliverables:</b> Verification records, field standards and revision lists, member or news pages, and a publication review workflow.",
        ],
    ),
    (
        "AI Agent-enabled Operations Workflow",
        [
            "<b>Context:</b> Academic-organization operations span public sources, member data, news content, page status and multiple milestones. Repeated tasks can drift when decision boundaries and acceptance criteria are unclear.",
            "<b>Actions:</b> Structures work into research, field organization, content generation, page checks and public-result review; sets source priorities, output formats, no-guess boundaries and acceptance criteria for AI agents, then iterates the rules when results drift.",
            "<b>Deliverables:</b> Structured verification records and revision lists, member and news content, source and page review results, and reusable skills, workflows and checklists.",
        ],
    ),
    (
        "Film Awards Materials & Milestone Coordination",
        [
            "<b>Context:</b> Awards programs connect nominations, film materials, director or jury communication, review milestones, certificates and documentation.",
            "<b>Actions:</b> Organizes nomination and film materials, records missing or unverified items, follows communication, checks certificate wording and archives final materials.",
            "<b>Deliverables:</b> Nomination and film materials, open-item records, communication follow-up, certificate text and archived records.",
        ],
    ),
]


def register_fonts():
    pdfmetrics.registerFont(TTFont("ResumeSans", str(FONT_DIR / "Arial Unicode.ttf")))
    pdfmetrics.registerFont(TTFont("ResumeSansBold", str(FONT_DIR / "Arial Unicode.ttf")))
    pdfmetrics.registerFontFamily("ResumeSans", normal="ResumeSans", bold="ResumeSansBold")


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("Name", parent=base["Normal"], fontName="ResumeSansBold", fontSize=25, leading=29, textColor=INK, spaceAfter=2),
        "role": ParagraphStyle("Role", parent=base["Normal"], fontName="ResumeSansBold", fontSize=10.2, leading=14, textColor=VIOLET, spaceAfter=6),
        "lead": ParagraphStyle("Lead", parent=base["Normal"], fontName="ResumeSans", fontSize=9.1, leading=14.2, textColor=MUTED),
        "contact": ParagraphStyle("Contact", parent=base["Normal"], fontName="ResumeSans", fontSize=8.1, leading=12.4, textColor=MUTED),
        "section": ParagraphStyle("Section", parent=base["Normal"], fontName="ResumeSansBold", fontSize=12.2, leading=15, textColor=INK, spaceBefore=6, spaceAfter=4),
        "job": ParagraphStyle("Job", parent=base["Normal"], fontName="ResumeSansBold", fontSize=9.5, leading=13.4, textColor=INK, spaceAfter=1),
        "meta": ParagraphStyle("Meta", parent=base["Normal"], fontName="ResumeSans", fontSize=8.0, leading=11.4, textColor=VIOLET, spaceAfter=2),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontName="ResumeSans", fontSize=8.2, leading=12.7, textColor=MUTED, alignment=TA_LEFT, spaceAfter=2.5),
        "bullet": ParagraphStyle("Bullet", parent=base["Normal"], fontName="ResumeSans", fontSize=8.15, leading=12.6, textColor=MUTED, leftIndent=10, firstLineIndent=-7, bulletIndent=2, spaceAfter=2),
        "small": ParagraphStyle("Small", parent=base["Normal"], fontName="ResumeSans", fontSize=7.2, leading=10.4, textColor=MUTED),
    }


def p(text, style):
    return Paragraph(text, style)


def pdf_bullets(items, style):
    return [p(f"• {item}", style) for item in items]


def section_title(text, s):
    table = Table([[p(text, s["section"])]], colWidths=[174 * mm])
    table.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 0.75, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return KeepTogether([Spacer(1, 2), table])


def keyword_bar(items, s):
    table = Table([[p(item, s["small"]) for item in items]], colWidths=[34.8 * mm] * 5)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), PALE),
        ("BOX", (0, 0), (-1, -1), 0.5, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return table


def skills_table(groups, s):
    rows = [[p(f"<b>{label}</b>", s["body"]), p(text, s["body"])] for label, text in groups]
    table = Table(rows, colWidths=[36 * mm, 138 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), PALE),
        ("BOX", (0, 0), (-1, -1), 0.5, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return table


def page_decor(label):
    def draw(canvas, doc):
        canvas.saveState()
        width, height = A4
        canvas.setFillColor(VIOLET)
        canvas.rect(0, height - 7 * mm, width, 7 * mm, stroke=0, fill=1)
        canvas.setFont("ResumeSans", 7)
        canvas.setFillColor(MUTED)
        canvas.drawString(18 * mm, 10 * mm, label)
        canvas.drawRightString(width - 18 * mm, 10 * mm, str(doc.page))
        canvas.restoreState()
    return draw


def pdf_document(path, title, footer):
    doc = BaseDocTemplate(str(path), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=15 * mm, bottomMargin=16 * mm, title=title, author="Eacon Jing")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="resume", frames=frame, onPage=page_decor(footer))])
    return doc


def add_pdf_header(story, s, name, role, lead, contact):
    story.extend([p(name, s["name"]), p(role, s["role"])])
    table = Table([[p(lead, s["lead"]), p(contact, s["contact"])]], colWidths=[122 * mm, 52 * mm])
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (0, 0), 6 * mm),
        ("LEFTPADDING", (1, 0), (1, 0), 2 * mm),
    ]))
    story.extend([table, Spacer(1, 4)])


def build_pdf(language, s):
    zh = language == "zh"
    story = []
    if zh:
        add_pdf_header(
            story,
            s,
            "江奕坤 / Eacon Jing",
            "AI 应用运营与交付",
            "求职方向：AI 应用运营、AI 项目交付运营、AI 工具运营、AI / SaaS 客户成功、AI 用户教育与应用支持、知识库运营与信息治理。",
            "hieacon@gmail.com<br/>WeChat: JovoeYK_05<br/>中国",
        )
        story += [section_title("职业定位", s), p("具备多年内容、教育与机构运营经验，能够理解业务需求、拆解任务、组织 Codex 与 ChatGPT 等 AI Agent 协作，并通过信息核验、跨角色推进和结果复核，把复杂事项推进到可用、可验证的交付。优势不是单纯会使用工具，而是能够定义目标与边界、持续纠偏并推动结果闭环。", s["body"])]
        story += [section_title("核心能力", s), keyword_bar(["需求分析", "任务拆解", "AI Agent 协作", "信息治理", "交付推进"], s)]
        story += [section_title("工作经历", s)]
        for name, meta, points in ZH_JOBS:
            story.append(KeepTogether([p(name, s["job"]), p(meta, s["meta"])] + pdf_bullets(points, s["bullet"])))
        story += [PageBreak(), section_title("代表项目", s)]
        for title, points in ZH_CASES:
            story.append(KeepTogether([p(title, s["job"])] + [p(point, s["body"]) for point in points]))
        story += [section_title("技能", s), skills_table([
            ("需求与交付", "需求分析、任务拆解、优先级、责任边界、验收标准与结果复核"),
            ("AI Agent 协作", "Codex / ChatGPT 协作、AI Agent Workflow、Vibe Coding、迭代纠偏"),
            ("信息与知识", "官方来源检索、多源核验、结构化信息整理、状态追踪、内容与知识库运营"),
            ("运营与标准化", "项目协同、反馈闭环、Skill、模板、台账、工作流与检查清单设计"),
        ], s)]
        story += [section_title("教育背景", s), p("伊犁师范大学 | 英语教育方向，本科 | 2012.09 - 2016.06", s["job"]), p("接受英语语言、教育学、课程设计和教学实践训练，具备教育内容组织与英文书面材料处理基础。", s["body"])]
        path = PUBLIC / "eacon-jing-resume.pdf"
        doc = pdf_document(path, "Eacon Jing Resume - Chinese", "Eacon Jing · AI 应用运营与交付")
    else:
        add_pdf_header(
            story,
            s,
            "Eacon Jing",
            "AI Application Operations & Delivery",
            "Target roles: AI Application Operations, AI Delivery Operations, AI Tool Operations, Customer Success (AI / SaaS), AI User Education & Enablement, and Knowledge Operations or Information Governance.",
            "hieacon@gmail.com<br/>WeChat: JovoeYK_05<br/>China",
        )
        story += [section_title("Professional Profile", s), p("Years of experience across content, education and organization operations. Translates business needs into executable work, coordinates Codex, ChatGPT and other AI agents, and uses information verification, cross-role follow-through and outcome review to move complex work toward usable, verifiable delivery. The differentiator is not tool access alone, but defining boundaries, correcting drift and closing the loop.", s["body"])]
        story += [section_title("Core Capabilities", s), keyword_bar(["Requirement Analysis", "Task Decomposition", "AI Agent Coordination", "Information Governance", "Delivery"], s)]
        story += [section_title("Experience", s)]
        for name, meta, points in EN_JOBS:
            story.append(KeepTogether([p(name, s["job"]), p(meta, s["meta"])] + pdf_bullets(points, s["bullet"])))
        story += [PageBreak(), section_title("Selected Projects", s)]
        for title, points in EN_CASES:
            story.append(KeepTogether([p(title, s["job"])] + [p(point, s["body"]) for point in points]))
        story += [section_title("Skills", s), skills_table([
            ("Requirements & Delivery", "Requirement analysis, task decomposition, priorities, ownership boundaries, acceptance criteria and outcome review"),
            ("AI Agent Coordination", "Codex / ChatGPT collaboration, AI Agent Workflow, vibe coding and iterative correction"),
            ("Information & Knowledge", "Official-source research, multi-source verification, structured information, status tracking, content and knowledge operations"),
            ("Operations & Standards", "Project coordination, feedback loops, reusable skills, templates, trackers, workflows and checklists"),
        ], s)]
        story += [section_title("Education", s), p("Yili Normal University | Undergraduate study in English Education | Sep 2012 - Jun 2016", s["job"]), p("Training in English language, pedagogy, curriculum design and teaching practice, with a foundation in education content and written English materials.", s["body"])]
        path = PUBLIC / "eacon-jing-resume-en.pdf"
        doc = pdf_document(path, "Eacon Jing Resume - English", "Eacon Jing · AI Application Operations & Delivery")
    doc.build(story)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


DOCX_FONT = "Noto Sans CJK SC"
DOCX_FONT_PATH = Path("/Library/Fonts/NotoSansCJKsc-Regular.otf")


def set_run_font(run, name=DOCX_FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_docx_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def embed_docx_font(doc, font_path=DOCX_FONT_PATH, font_name=DOCX_FONT):
    font = FontToolsTTFont(font_path)
    font.flavor = None
    font_bytes_path = ROOT / "tmp" / "embedded-docx-font.ttf"
    font_bytes_path.parent.mkdir(parents=True, exist_ok=True)
    font.save(font_bytes_path)
    font_bytes = font_bytes_path.read_bytes()

    font_part = Part(PackURI("/word/fonts/noto-sans-cjk-sc.otf"), "application/x-font-opentype", font_bytes, doc.part.package)
    fonts_part = doc.part.part_related_by(RT.FONT_TABLE)
    rel_id = fonts_part.relate_to(
        font_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
    )
    fonts_el = parse_xml(fonts_part.blob)
    font_el = OxmlElement("w:font")
    font_el.set(qn("w:name"), font_name)
    embed = OxmlElement("w:embedRegular")
    embed.set(qn("r:id"), rel_id)
    font_el.append(embed)
    fonts_el.append(font_el)
    fonts_part._blob = etree.tostring(fonts_el, xml_declaration=True, encoding="UTF-8", standalone="yes")


def build_docx_resume():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOCX_FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT)
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(2.8)
    normal.paragraph_format.line_spacing = 1.06
    for style_name, size, color, before, after in (
        ("Heading 1", 13.2, (17, 17, 17), 6, 3),
        ("Heading 2", 10.3, (111, 75, 232), 3.5, 1.5),
    ):
        style = styles[style_name]
        style.font.name = DOCX_FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = styles["List Bullet"]
    bullet.font.name = DOCX_FONT
    bullet._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT)
    bullet.font.size = Pt(9.0)
    bullet.paragraph_format.space_after = Pt(1.8)
    bullet.paragraph_format.line_spacing = 1.04

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    set_run_font(title.add_run("江奕坤 / Eacon Jing"), size=24, bold=True, color=(17, 17, 17))
    role = doc.add_paragraph()
    role.paragraph_format.space_after = Pt(4)
    set_run_font(role.add_run("AI 应用运营与交付"), size=11, bold=True, color=(111, 75, 232))
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact.paragraph_format.space_after = Pt(4)
    set_run_font(contact.add_run("hieacon@gmail.com  |  WeChat: JovoeYK_05  |  中国"), size=8.6, color=(80, 87, 99))

    doc.add_heading("职业定位", level=1)
    doc.add_paragraph("求职方向：AI 应用运营、AI 项目交付运营、AI 工具运营、AI / SaaS 客户成功、AI 用户教育与应用支持、知识库运营与信息治理。具备多年内容、教育与机构运营经验，能够理解业务需求、拆解任务、组织 Codex 与 ChatGPT 等 AI Agent 协作，并通过信息核验、跨角色推进和结果复核，把复杂事项推进到可用、可验证的交付。")

    doc.add_heading("核心能力", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    for cell, text in zip(table.rows[0].cells, ["需求分析", "任务拆解", "AI Agent 协作", "信息治理", "交付推进"]):
        cell.width = Inches(1.28)
        set_cell_margins(cell)
        cell.vertical_alignment = 1
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cell.paragraphs[0].add_run(text), size=8.2, bold=True, color=(80, 87, 99))

    doc.add_heading("工作经历", level=1)
    for name, meta, points in ZH_JOBS:
        doc.add_heading(name, level=2)
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(2)
        set_run_font(meta_p.add_run(meta), size=8.5, color=(111, 75, 232))
        for point in points:
            add_docx_bullet(doc, point)

    doc.add_heading("代表项目", level=1)
    for title_text, points in ZH_CASES:
        doc.add_heading(title_text, level=2)
        for point in points:
            clean = point.replace("<b>", "").replace("</b>", "")
            doc.add_paragraph(clean)

    doc.add_heading("技能", level=1)
    skill_groups = [
        ("需求与交付", "需求分析、任务拆解、优先级、责任边界、验收标准与结果复核"),
        ("AI Agent 协作", "Codex / ChatGPT 协作、AI Agent Workflow、Vibe Coding、迭代纠偏"),
        ("信息与知识", "官方来源检索、多源核验、结构化信息整理、状态追踪、内容与知识库运营"),
        ("运营与标准化", "项目协同、反馈闭环、Skill、模板、台账、工作流与检查清单设计"),
    ]
    table = doc.add_table(rows=len(skill_groups), cols=2)
    table.autofit = False
    for row, (label, value) in zip(table.rows, skill_groups):
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(5.2)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = 1
        set_run_font(row.cells[0].paragraphs[0].add_run(label), size=8.8, bold=True, color=(17, 17, 17))
        set_run_font(row.cells[1].paragraphs[0].add_run(value), size=8.8, color=(80, 87, 99))

    doc.add_heading("教育背景", level=1)
    doc.add_heading("伊犁师范大学 | 英语教育方向，本科 | 2012.09 - 2016.06", level=2)
    doc.add_paragraph("接受英语语言、教育学、课程设计和教学实践训练，具备教育内容组织与英文书面材料处理基础。")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.name is None:
                            set_run_font(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Eacon Jing | AI 应用运营与交付"), size=7.5, color=(120, 124, 132))

    embed_docx_font(doc)
    path = ROOT / "江奕坤EaconJing_AI应用运营与交付_简历.docx"
    doc.save(path)
    public_path = PUBLIC / "eacon-jing-resume.docx"
    public_path.write_bytes(path.read_bytes())


if __name__ == "__main__":
    register_fonts()
    styles = pdf_styles()
    build_pdf("zh", styles)
    build_pdf("en", styles)
    build_docx_resume()
    print("Built Chinese and English PDF resumes plus editable Chinese DOCX resume")
