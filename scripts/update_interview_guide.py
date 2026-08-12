from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "江奕坤EaconJing_官网简历面试解释手册.docx"


PARAGRAPH_REPLACEMENTS = {
    "涉及 AI / vibe coding 时，定位为“AI 辅助的运营交付能力”，不要把自己包装成纯技术工程师。":
        "涉及 AI Agent 时，定位为“人机协同的运营交付能力”，不要把自己包装成纯技术工程师。",
    "我最早是英语教育背景，做过一线英语教学，后来转向教育内容和新媒体运营，再到教育科技机构的管理岗位。现在我参与 NAAI 及相关国际学术与文化项目运营，工作内容包括会员拓展与资料审核、官网内容维护、国际合作邮件沟通、证书和奖项流程支持，以及用 AI 工具提升双语文档、邮件模板和官网原型的交付效率。":
        "我最早是英语教育背景，做过一线英语教学，后来转向教育内容和新媒体运营，再到教育科技机构的管理岗位。现在我参与 NAAI 及相关国际学术与文化项目运营，工作内容包括会员拓展与资料审核、官网内容维护、国际合作邮件沟通、证书和奖项流程支持，以及设计人机协同 AI Agent 工作流，提升双语文档、邮件模板和官网原型的交付效率。",
    "我做这个官网，一方面是为了把经历从传统简历里“压缩的几行字”展开成更清楚的能力结构；另一方面也是为了展示我现在的工作方式。我并不只是说自己会 AI 工具，而是把 AI 辅助内容、页面原型、双语表达和信息组织实际做成一个可访问的作品。它既是个人品牌页，也是一个小型运营交付样本。":
        "我做这个官网，一方面是为了把经历从传统简历里“压缩的几行字”展开成更清楚的能力结构；另一方面也是为了展示我现在的工作方式。我不只是泛泛地说自己会使用 AI，而是把需求拆解、资料整理、双语内容、页面实现、人工复核和上线检查组织成完整的 AI Agent 工作流。它既是个人品牌页，也是一个小型运营交付样本。",
    "不要说成：“这个网站完全证明我会开发。”更稳的说法是：“它证明我能用 AI 工具和现有技术工作流，把个人信息、视觉表达、双语内容和表单流程整合成一个可上线的成果。”":
        "不要说成：“这个网站完全证明我会开发。”更稳的说法是：“它证明我能设计并推进人机协同 AI Agent 工作流，把个人信息、视觉表达、双语内容和表单流程整合成一个可上线、可检查的成果。”",
    "What makes my profile valuable is the combination of bilingual institutional communication, structured project operations, and AI-assisted workflow. I am not positioning myself as a pure software engineer. Instead, I use AI tools and vibe coding workflows to accelerate operational deliverables, such as website prototypes, bilingual documents, email templates, process drafts, and project materials.":
        "What makes my profile valuable is the combination of bilingual institutional communication, structured project operations and human-in-the-loop AI Agent workflows. I am not positioning myself as a pure software engineer. I decompose operational needs into clear steps, use AI Agents for suitable tasks, and remain responsible for review, implementation and final delivery.",
    "AI 工具与 Vibe Coding": "AI Agent 工作流与人机协同",
    "标准回答：我的定位是 AI 辅助运营交付，不是把自己包装成纯技术开发。我用 AI 帮助快速产出官网原型、双语文档、邮件模板、流程草案和项目材料，再由我做判断、修改和落地。":
        "标准回答：我的定位是 AI Agent 辅助的运营交付，不是把自己包装成纯技术开发。我会先拆解目标和流程，再让 AI Agent 协助资料整理、结构化和初稿生成，最后由我核对事实、调整表达、确认实现并完成交付检查。",
    "答：我不是传统意义上的全职程序员，但我能使用 AI 和 vibe coding 工作流，把需求拆清楚、生成原型、检查效果、不断修改，最终交付可用页面或工具。":
        "答：我不是传统意义上的全职程序员，但我能完成任务拆解与流程编排，使用 AI Agent 和现有工具生成原型，并通过人工复核、效果检查和持续修改，最终交付可用页面或运营工具。",
    "适配场景：如果岗位重视 AI 工具、数字化、低成本搭建、内容效率，要重点讲这一块。":
        "适配场景：如果岗位重视 AI Agent、数字化、低成本搭建和内容效率，要重点讲任务拆解、人机协同与交付检查。",
    "答：我不是纯技术工程师，我更准确的定位是懂 AI 工具和低代码/AI 辅助工作流的运营型人才。我的优势是把业务需求、内容表达和工具交付连接起来。":
        "答：我不是纯技术工程师，我更准确的定位是能设计 AI Agent 工作流的运营型人才。我的优势是把业务目标拆成清晰步骤，让 AI Agent 处理适合自动化的环节，再由人工负责事实、语气、边界和最终交付。",
    "答：复合型。很多人只懂内容或只懂执行，我的优势是语言、内容、项目流程、机构运营和 AI 工具能结合。尤其适合需要对外沟通、材料交付和跨团队推进的岗位。":
        "答：复合型。很多人只懂内容或只懂执行，我的优势是语言、内容、项目流程、机构运营和 AI Agent 工作流能结合。尤其适合需要对外沟通、材料交付和跨团队推进的岗位。",
    "问：AI / vibe coding 会不会只是噱头？": "问：AI Agent 会不会只是噱头？",
    "我的优势是复合型：教育理解、内容表达、机构运营、国际沟通和 AI 工具能结合。":
        "我的优势是复合型：教育理解、内容表达、机构运营、国际沟通和 AI Agent 工作流能结合。",
    "附录：16 个技能逐项解释": "附录：16 个技能逐项解释",
    "AI 交付类：AI 工具应用、Vibe Coding、AI-assisted Workflow、官网原型搭建。":
        "AI 交付类：AI Agent 工作流设计、Human-in-the-loop 人机协同、任务拆解与流程编排、官网原型搭建。",
    "AI 工具应用": "AI Agent 工作流设计",
    "Vibe Coding": "任务拆解与流程编排",
    "AI-assisted Workflow": "Human-in-the-loop 人机协同",
}


CELL_REPLACEMENTS = {
    "3": "4",
    "官网公开奖项类别，说明奖项体系存在多个方向，我参与相关流程支持。":
        "官网公开提名类奖项数量，用于说明项目环境包含多个奖项方向，我参与相关运营流程支持。",
    "20": "17",
    "公开官网页面数量，说明官网内容运营有一定复杂度，我参与内容建设与维护。":
        "官网公开导航栏目数量，说明机构信息分布在多个栏目，我参与相关内容建设与维护。",
    "AI 工具应用、Vibe Coding、AI-assisted Workflow、官网原型搭建":
        "AI Agent 工作流设计、Human-in-the-loop 人机协同、任务拆解与流程编排、官网原型搭建",
    "我能把 AI 当作效率工具，把需求转成可见成果。":
        "我能把任务拆成可执行步骤，让 AI Agent 与人工复核协同完成可检查的成果。",
    "AI 辅助原型、流程文档、邮件模板、官网页面、材料整理":
        "AI Agent 流程、任务拆解、人工复核、官网页面、文档和材料整理",
    "我能用 AI 和 vibe coding 工作流完成运营型页面原型和可上线成果。":
        "我能通过任务拆解、人机协同和交付检查完成运营型页面原型与可上线成果。",
    "“我能用 AI 和 vibe coding 工作流完成运营型页面原型和可上线成果。”":
        "“我能通过任务拆解、人机协同和交付检查完成运营型页面原型与可上线成果。”",
    "“20 个页面全部由我独立开发。”": "“17 个公开栏目全部由我独立开发。”",
}


SKILL_CELL_REPLACEMENTS = {
    "把 AI 用在信息整理、初稿生成、文案优化、表格梳理、邮件草稿、页面原型和流程设计中。":
        "围绕明确目标，把资料整理、结构化、初稿生成、人工复核和交付检查组织成可重复的 AI Agent 工作流。",
    "我把 AI 当协作工具，不把它当最终答案。AI 提速，我负责判断、核对事实、调整语气和落地。":
        "我会先定义输入、输出和检查标准，再安排 AI Agent 处理适合自动化的步骤，由我负责判断、事实核对和最终落地。",
    "用 AI 辅助生成官网结构、双语资料、邮件模板、项目流程草案和面试准备文档。":
        "用 AI Agent 工作流完成官网结构、双语资料、邮件模板、项目流程草案和面试准备文档。",
    "问“怎么避免 AI 错？”答：所有事实、数字、身份、链接、时间都人工核对，AI 只做初稿和整理。":
        "问“怎么避免 AI 错？”答：先限定资料来源和任务边界，所有事实、数字、身份、链接与时间仍由人工核对。",
    "不要说“AI 能替我完成所有工作”。":
        "不要说“AI 能替我完成所有判断和交付”。",
    "用自然语言描述需求，借助 AI 和代码工具快速做出页面、原型、小工具或自动化流程，再不断检查和修改。":
        "把模糊需求拆成目标、输入、步骤、责任边界、检查点和交付物，再安排合适的 AI Agent 或工具逐步执行。",
    "我不是把自己包装成传统程序员，而是能用 vibe coding 把运营需求快速转成可用的数字化成果。":
        "我擅长把复杂运营需求拆清楚，并编排资料整理、内容生成、人工确认、页面实现和最终检查。",
    "个人官网的页面结构、中英切换、表单入口、部署上线，就是一个适合说明的例子。":
        "个人官网从双语内容、表单入口、PDF 简历到部署上线，是任务拆解与流程编排的完整例子。",
    "问“你会开发吗？”答：我不是全栈工程师，但能用 AI 辅助工作流完成运营型网站和工具原型。":
        "问“你会开发吗？”答：我不是全栈工程师，但能拆解业务需求、组织工具链，并完成运营型网站和工具原型。",
    "不要直接说“我是全栈开发”。":
        "不要把任务编排能力夸大成复杂底层系统开发能力。",
    "把 AI 嵌入日常工作流程，从资料输入、信息整理、草案生成、人工校对到最终交付形成闭环。":
        "明确哪些环节适合 AI Agent，哪些环节必须由人确认，让效率提升和责任判断同时存在。",
    "我的 AI 工作流通常是：先定义目标和材料，再生成初稿或结构，然后人工判断、修改、验证，最后沉淀成模板。":
        "我的流程通常是：AI Agent 协助研究、整理和初稿；人负责来源判断、事实核对、语气、边界、实现验收和最终签字。",
    "官网、双语文档、邮件模板、项目说明、流程表、面试手册都可以按这个流程产出。":
        "会员资料、官网、双语文档、邮件模板、项目说明和面试手册都采用人机协同流程。",
    "问“和普通使用 ChatGPT 有什么区别？”答：区别在于是否形成稳定流程和可复用交付物，而不是临时问一句。":
        "问“人为什么还要参与？”答：机构项目涉及身份、事实、语气和公开责任，这些必须由人确认，不能直接接受模型输出。",
}


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_paragraphs(doc):
    changed = 0
    for paragraph in doc.paragraphs:
        replacement = PARAGRAPH_REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            set_paragraph_text(paragraph, replacement)
            changed += 1
    return changed


def replace_table_cells(doc):
    changed = 0
    replacements = {**CELL_REPLACEMENTS, **SKILL_CELL_REPLACEMENTS}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replacement = replacements.get(cell.text)
                if replacement is not None:
                    set_paragraph_text(cell.paragraphs[0], replacement)
                    for paragraph in cell.paragraphs[1:]:
                        set_paragraph_text(paragraph, "")
                    changed += 1
    return changed


def copy_run_format(source, target):
    target._r.get_or_add_rPr()
    if source._r.rPr is not None:
        target._r.rPr = deepcopy(source._r.rPr)


def insert_paragraph(anchor, text, style="Normal", page_break_before=False):
    paragraph = anchor.insert_paragraph_before(style=style)
    run = paragraph.add_run(text)
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.keep_with_next = style.startswith("Heading")
    return paragraph


def add_case_questions(doc):
    if any(p.text == "十六、两份官网案例的完整面试问答" for p in doc.paragraphs):
        return
    anchor = next((p for p in doc.paragraphs if p.text == "附录：16 个技能逐项解释"), None)
    if anchor is None:
        raise RuntimeError("Appendix anchor not found")

    insert_paragraph(anchor, "十六、两份官网案例的完整面试问答", "Heading 1", page_break_before=True)
    insert_paragraph(anchor, "案例一：国际学术会员体系运营", "Heading 2")
    insert_paragraph(
        anchor,
        "30 秒讲法：会员资料横跨不同国家、学科和职业背景，我参与的重点是把身份、机构、领域、公开来源和待补信息整理成可审核、可沟通、可维护的结构。AI Agent 协助整理材料和生成清单，最终由人工核对姓名、头衔、机构、链接和术语。",
    )
    insert_paragraph(anchor, "问：你个人到底做了什么？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：我承担的是具体运营动作，包括资料收集与结构化、公开身份核对、待补信息标记、补充沟通、双语初稿和会员页面维护支持。我不会把机构整体会员规模说成个人独立拓展成果。",
    )
    insert_paragraph(anchor, "问：AI Agent 在这个案例里做什么？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：它适合做材料归类、字段提取、缺失项识别、核对清单和双语初稿。涉及身份真实性、头衔、机构归属、来源可信度和最终公开内容时，必须由人复核。",
    )
    insert_paragraph(anchor, "问：如何解释 3,200+ 和 85+？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：这是 NAAI 官网公开呈现的机构背景，用来说明会员体系的规模和跨国复杂度。我的个人贡献只落在我实际参与的资料、审核支持、沟通和维护环节。",
    )

    insert_paragraph(anchor, "案例二：AI Agent 官网与运营工作流", "Heading 2")
    insert_paragraph(
        anchor,
        "30 秒讲法：我把个人官网当成一份真实的运营交付。先定义双语简历、项目案例、表单访问和 PDF 下载目标，再拆成资料、内容、页面、表单、人工复核、测试和部署七类任务，最终形成中英文一致、手机可用、能够上线的成果。",
    )
    insert_paragraph(anchor, "问：这和普通使用 ChatGPT 有什么不同？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：普通使用可能只是临时生成一段文字；AI Agent 工作流有明确输入、步骤、检查点和交付标准。这个项目还包含真实表单提交、双语 PDF、手机适配、浏览器测试和上线验证，不是生成文字后就结束。",
    )
    insert_paragraph(anchor, "问：你在技术上能做到什么程度？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：我能把业务需求拆解清楚，借助 AI Agent 和现有工具完成运营型网站原型、内容结构、表单流程、双语版本、测试和部署。我不把自己包装成全栈工程师，也不声称独立开发复杂底层系统。",
    )
    insert_paragraph(anchor, "问：Human-in-the-loop 体现在哪里？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：模型负责提高研究、整理、初稿和实现速度；我负责核对数字与来源、控制中英文口径、确认页面效果、测试交互、判断边界并决定最终发布。人始终对公开成果负责。",
    )
    insert_paragraph(anchor, "问：你的标准工作方法是什么？", "Heading 3")
    insert_paragraph(
        anchor,
        "答：六步：明确目标、核对资料、结构化拆解、AI 辅助初稿、人工复核、交付检查。面试时可以把任何一个官网、邮件、会员或奖项案例都按这六步展开。",
    )


def apply_cjk_font(doc):
    font_name = "Noto Sans CJK SC"

    def set_run_font(run):
        run.font.name = font_name
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.get_or_add_rFonts()
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:eastAsia"), font_name)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)


def validate(doc):
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    forbidden = ["Vibe Coding", "vibe coding", "AI 工具应用", "AI-assisted Workflow"]
    found = [term for term in forbidden if term in text]
    if found:
        raise RuntimeError(f"Outdated terms remain: {found}")
    required = ["4", "17", "24", "AI Agent 工作流", "Human-in-the-loop", "任务拆解与流程编排"]
    missing = [term for term in required if term not in text]
    if missing:
        raise RuntimeError(f"Required terms missing: {missing}")


def main():
    doc = Document(DOCX_PATH)
    paragraph_changes = replace_paragraphs(doc)
    cell_changes = replace_table_cells(doc)
    add_case_questions(doc)
    apply_cjk_font(doc)
    validate(doc)
    output = DOCX_PATH.with_name(f"{DOCX_PATH.stem}.updated.docx")
    doc.save(output)
    output.replace(DOCX_PATH)
    print(f"Updated {DOCX_PATH.name}: {paragraph_changes} paragraphs, {cell_changes} table cells")


if __name__ == "__main__":
    main()
